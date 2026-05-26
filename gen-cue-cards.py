#!/usr/bin/env python3
"""
Gera cue cards (DOCX) a partir do palestra-bruno-pinha.html
A4 portrait, 2x2 cartões por página = 6 páginas, 24 cartões total.
"""
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 1. Parse data-notes do HTML
with open('/root/palestra/palestra-bruno-pinha.html', 'r', encoding='utf-8') as f:
    soup = BeautifulSoup(f.read(), 'html.parser')

slides_data = []
for slide in soup.find_all('div', class_='slide'):
    sid = slide.get('id', '')
    notes = slide.get('data-notes', '').strip()
    slides_data.append({'id': sid, 'notes': notes})

# 2. Metadata por slide (título + bloco + horário do roteiro)
SLIDE_META = {
    's1':  {'num': '01', 'title': 'Capa',                                      'block': 'Bloco 1 · Abertura provocativa',  'time': '21h00 – 21h01'},
    's2':  {'num': '02', 'title': 'O que conecta tudo isso?',                  'block': 'Bloco 1 · Abertura provocativa',  'time': '21h01 – 21h03'},
    's3':  {'num': '03', 'title': 'Frase-tese',                                'block': 'Bloco 1 · Abertura provocativa',  'time': '21h03 – 21h05'},
    's4':  {'num': '04', 'title': 'Carreira como stack profissional',          'block': 'Bloco 2 · A carreira como stack', 'time': '21h05 – 21h06'},
    's5':  {'num': '05', 'title': 'Camada 01 — Eletrônica e mundo físico',     'block': 'Bloco 2 · A carreira como stack', 'time': '21h06 – 21h07'},
    's6':  {'num': '06', 'title': 'Camada 02 — Delphi & Legados',              'block': 'Bloco 2 · A carreira como stack', 'time': '21h07 – 21h08'},
    's7':  {'num': '07', 'title': 'Camada 03 — Arquitetura e legados',         'block': 'Bloco 2 · A carreira como stack', 'time': '21h08 – 21h10'},
    's8':  {'num': '08', 'title': 'Camada 04 — Gestão de projetos',            'block': 'Bloco 2 · A carreira como stack', 'time': '21h10 – 21h11'},
    's9':  {'num': '09', 'title': 'Camada 05 — Empreendedorismo (Conmarket)',  'block': 'Bloco 2 · A carreira como stack', 'time': '21h11 – 21h13'},
    's10': {'num': '10', 'title': 'Camada 06 — Mestrado UEM',                  'block': 'Bloco 2 · A carreira como stack', 'time': '21h13 – 21h14'},
    's11': {'num': '11', 'title': 'Camada 07 — PesoCam (pico)',                'block': 'Bloco 2 · A carreira como stack', 'time': '21h14 – 21h15'},
    's12': {'num': '12', 'title': 'O problema — peso na pecuária',             'block': 'Bloco 3 · O problema real',       'time': '21h15 – 21h18'},
    's13': {'num': '13', 'title': 'A hipótese — sem contato físico?',          'block': 'Bloco 3 · O problema real',       'time': '21h18 – 21h20'},
    's14': {'num': '14', 'title': 'Pipeline — câmera → borda → IA → peso',     'block': 'Bloco 3 · O problema real',       'time': '21h20 – 21h22'},
    's15': {'num': '15', 'title': 'Por que é difícil — 6 fricções do campo',   'block': 'Bloco 3 · O problema real',       'time': '21h22 – 21h25'},
    's16': {'num': '16', 'title': '9 domínios — engenharia conecta tudo',      'block': 'Bloco 4 · Engenharia da inovação','time': '21h25 – 21h30'},
    's17': {'num': '17', 'title': 'A IA estima. A engenharia confia.',         'block': 'Bloco 4 · Engenharia da inovação','time': '21h30 – 21h35'},
    's18': {'num': '18', 'title': 'Validações — GDIN, Korean Valley, UEM',     'block': 'Bloco 4 · Engenharia da inovação','time': '21h35 – 21h40'},
    's19': {'num': '19', 'title': 'IA amplifica entendimento (não substitui)', 'block': 'Bloco 5 · IA e responsabilidade', 'time': '21h40 – 21h44'},
    's20': {'num': '20', 'title': '4 princípios não-negociáveis',              'block': 'Bloco 5 · IA e responsabilidade', 'time': '21h44 – 21h48'},
    's21': {'num': '21', 'title': '4 pilares — pra vocês, agora',              'block': 'Bloco 6 · Carreira e networking', 'time': '21h48 – 21h51'},
    's22': {'num': '22', 'title': 'Networking — confiança antes da oportunidade', 'block': 'Bloco 6 · Carreira e networking', 'time': '21h51 – 21h55'},
    's23': {'num': '23', 'title': 'Docência — próxima camada (Unicesumar)',    'block': 'Bloco 7 · Fechamento',            'time': '21h55 – 21h58'},
    's24': {'num': '24', 'title': 'Fechamento + contatos + QR',                'block': 'Bloco 7 · Fechamento',            'time': '21h58 – 22h00'},
}

slides = []
for s in slides_data:
    meta = SLIDE_META.get(s['id'], {'num': '??', 'title': s['id'], 'block': '', 'time': ''})
    slides.append({**meta, 'notes': s['notes']})

print(f"{len(slides)} slides extraídos do HTML.")

# 3. Helpers para XML do python-docx
def set_cell_border(cell, color='888888', style='dashed', size='6'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove bordas anteriores se houver
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), style)
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=140, left=140, bottom=140, right=140):
    """Margins em twips (1/20 pt). 140 ≈ 2.5mm"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        m = OxmlElement(f'w:{name}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

# 4. Build doc
doc = Document()
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(6)
section.bottom_margin = Mm(6)
section.left_margin = Mm(6)
section.right_margin = Mm(6)

# Cor da identidade: teal #249F9C e dark #091817
TEAL = RGBColor(0x24, 0x9F, 0x9C)
DARK = RGBColor(0x09, 0x18, 0x17)
GRAY = RGBColor(0x88, 0x88, 0x88)
TEXT = RGBColor(0x22, 0x22, 0x22)

CARD_W_MM = 99   # ~ (210 - 6 - 6) / 2 = 99mm
CARD_H_MM = 94   # ~ (297 - 6 - 6) / 3 = 95mm — 3 linhas por página

CARDS_PER_PAGE = 6
TOTAL_PAGES = (len(slides) + CARDS_PER_PAGE - 1) // CARDS_PER_PAGE

for page_idx in range(TOTAL_PAGES):
    table = doc.add_table(rows=3, cols=2)
    table.autofit = False

    for col in table.columns:
        for cell in col.cells:
            cell.width = Mm(CARD_W_MM)

    for row in table.rows:
        row.height = Mm(CARD_H_MM)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    for r in range(3):
        for c in range(2):
            slide_idx = page_idx * CARDS_PER_PAGE + r * 2 + c
            if slide_idx >= len(slides):
                continue
            slide = slides[slide_idx]
            cell = table.rows[r].cells[c]
            set_cell_border(cell)
            set_cell_margins(cell, top=140, left=160, bottom=140, right=160)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            cell.text = ''

            # Linha 1: SLIDE NN  ·  21hXX – 21hYY
            p1 = cell.paragraphs[0]
            p1.paragraph_format.space_after = Pt(1)
            r1a = p1.add_run(f"SLIDE {slide['num']}")
            r1a.bold = True
            r1a.font.size = Pt(11)
            r1a.font.name = 'Arial'
            r1a.font.color.rgb = TEAL
            r1b = p1.add_run(f"   ·   {slide['time']}")
            r1b.font.size = Pt(8)
            r1b.font.name = 'Arial'
            r1b.font.color.rgb = GRAY

            # Linha 2: bloco
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            r2 = p2.add_run(slide['block'].upper())
            r2.font.size = Pt(6.5)
            r2.font.name = 'Arial'
            r2.font.color.rgb = GRAY

            # Linha 3: título do slide
            p3 = cell.add_paragraph()
            p3.paragraph_format.space_after = Pt(4)
            r3 = p3.add_run(slide['title'])
            r3.bold = True
            r3.font.size = Pt(9.5)
            r3.font.name = 'Arial'
            r3.font.color.rgb = DARK

            # Linha 4: notas
            p4 = cell.add_paragraph()
            p4.paragraph_format.space_after = Pt(0)
            p4.paragraph_format.line_spacing = 1.12
            r4 = p4.add_run(slide['notes'])
            r4.font.size = Pt(8.5)
            r4.font.name = 'Calibri'
            r4.font.color.rgb = TEXT

    if page_idx < TOTAL_PAGES - 1:
        doc.add_page_break()

OUT = '/root/palestra/cue-cards-bruno-pinha.docx'
doc.save(OUT)
print(f"DOCX salvo em {OUT}")
